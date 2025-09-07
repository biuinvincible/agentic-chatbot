"""
Core file processing functions for the Agentic Chatbot.
"""
import os
import re
import PyPDF2
import docx
import pandas as pd
from typing import List, Dict, Any, Optional
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_ollama import OllamaEmbeddings
# Import the new Docling processor
from agents.file_processing.docling_processor import extract_text_from_pdf_docling, extract_text_from_docx_docling

# Import the document classifier
import sys
import os
sys.path.append(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))
from classify_document import classify_document_for_processing

# Initialize the Ollama embedding model
print("Initializing Ollama embedding model...")
embedding_model = OllamaEmbeddings(model="bge-m3:567m")
print("Ollama embedding model initialized")

def extract_text_from_pdf(file_path: str, use_docling: bool = True) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        file_path (str): Path to the PDF file.
        use_docling (bool): Whether to use Docling VLM pipeline (default: True)
        
    Returns:
        str: Extracted text content.
    """
    if use_docling:
        # Use intelligent document classification to determine processing strategy
        from classify_document import classify_document_for_processing
        classification = classify_document_for_processing(file_path)
        
        # Use VLM processing only if recommended
        if classification.get("use_vlm", True):
            # Use Docling VLM pipeline with remote model
            return extract_text_from_pdf_docling(file_path)
        else:
            # Use fast text extraction for text-only documents
            pass  # Fall through to PyPDF2 extraction
    
    # Fallback to PyPDF2 for fast processing
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text.strip():  # Only add non-empty pages
                    text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            return text
    except Exception as e:
        return f"Error extracting text from PDF: {str(e)}"

def extract_text_from_docx(file_path: str, use_docling: bool = True) -> str:
    """
    Extract text from a Word document (.docx).
    
    Args:
        file_path (str): Path to the .docx file.
        use_docling (bool): Whether to use Docling VLM pipeline (default: True)
        
    Returns:
        str: Extracted text content.
    """
    if use_docling:
        # Use Docling VLM pipeline with remote model
        return extract_text_from_docx_docling(file_path)
    
    # Fallback to python-docx
    try:
        doc = docx.Document(file_path)
        text = ""
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text += f"{paragraph.text}\n"
        return text
    except Exception as e:
        return f"Error extracting text from DOCX: {str(e)}"

def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from a plain text file.
    
    Args:
        file_path (str): Path to the text file.
        
    Returns:
        str: Extracted text content.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        return f"Error reading text file: {str(e)}"

def extract_text_from_csv(file_path: str) -> str:
    """
    Extract and describe content from a CSV file.
    
    Args:
        file_path (str): Path to the CSV file.
        
    Returns:
        str: Description of the CSV content.
    """
    try:
        df = pd.read_csv(file_path)
        description = f"CSV file with {len(df)} rows and {len(df.columns)} columns.\n"
        description += f"Column names: {', '.join(df.columns)}\n"
        description += f"First few rows:\n{df.head().to_string()}"
        return description
    except Exception as e:
        return f"Error reading CSV file: {str(e)}"

def extract_text_from_xlsx(file_path: str) -> str:
    """
    Extract and describe content from an Excel file.
    
    Args:
        file_path (str): Path to the Excel file.
        
    Returns:
        str: Description of the Excel content.
    """
    try:
        xls = pd.ExcelFile(file_path)
        description = f"Excel file with {len(xls.sheet_names)} sheets: {', '.join(xls.sheet_names)}\n\n"
        
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name)
            description += f"Sheet '{sheet_name}': {len(df)} rows and {len(df.columns)} columns.\n"
            description += f"Column names: {', '.join(df.columns)}\n"
            description += f"First few rows:\n{df.head().to_string()}\n\n"
        return description
    except Exception as e:
        return f"Error reading Excel file: {str(e)}"

def extract_text_from_code(file_path: str) -> str:
    """
    Extract text from a code file.
    
    Args:
        file_path (str): Path to the code file.
        
    Returns:
        str: Extracted code content.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        return f"Error reading code file: {str(e)}"


def extract_text_from_image(file_path: str) -> str:
    """
    Handle image files by creating a placeholder description and storing file path.
    In a full implementation, this would use computer vision models.
    
    Args:
        file_path (str): Path to the image file.
        
    Returns:
        str: Description of the image file with file path information.
    """
    file_name = os.path.basename(file_path)
    _, ext = os.path.splitext(file_name)
    ext = ext.lower()
    
    return f"[Image File: {file_name}]\n" \
           f"This is an image file with extension {ext}. " \
           f"File path: {file_path}. " \
           f"To analyze this image, please use the image analysis agent. " \
           f"The image has been uploaded and can be referenced in your queries."


def chunk_text(text: str, chunk_size: int = 4000, chunk_overlap: int = 400) -> List[str]:
    """
    Split text into chunks using a hybrid approach optimized for large context windows.
    This approach combines semantic and structural chunking for better retrieval.
    
    Args:
        text (str): Text to chunk.
        chunk_size (int): Size of each chunk (larger for models with wide context windows).
        chunk_overlap (int): Overlap between chunks.
        
    Returns:
        List[str]: List of text chunks.
    """
    # First, split by sections (using double newlines as section separators)
    sections = re.split(r'\n\s*\n', text)
    
    # Initialize the text splitter for smaller chunks within sections
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len,
        separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""]
    )
    
    # Initialize list to hold all chunks
    all_chunks = []
    
    # Process each section
    for section in sections:
        # Skip empty sections
        if not section.strip():
            continue
            
        # If the section is smaller than the chunk size, add it as is
        if len(section) <= chunk_size:
            all_chunks.append(section.strip())
        else:
            # Split the section into smaller chunks
            section_chunks = text_splitter.split_text(section)
            all_chunks.extend(section_chunks)
    
    # Filter out empty chunks
    all_chunks = [chunk for chunk in all_chunks if chunk.strip()]
    
    return all_chunks

def create_vector_store(chunks: List[str], file_name: str, session_id: str) -> Chroma:
    """
    Create a vector store from text chunks, using session-specific collection.
    
    Args:
        chunks (List[str]): List of text chunks.
        file_name (str): Name of the file.
        session_id (str): Session ID for collection naming.
        
    Returns:
        Chroma: Vector store with the document chunks.
    """
    print(f"Creating vector store for file: {file_name} with session ID: {session_id}")
    print(f"Number of chunks: {len(chunks)}")
    
    # Create Document objects with metadata
    documents = []
    for i, chunk in enumerate(chunks):
        # Clean up the chunk text
        cleaned_chunk = chunk.strip()
        if cleaned_chunk:  # Only add non-empty chunks
            documents.append(
                Document(
                    page_content=cleaned_chunk,
                    metadata={
                        "source": file_name,
                        "session_id": session_id,
                        "chunk_index": i,
                        "total_chunks": len(chunks)
                    }
                )
            )
    
    print(f"Number of non-empty chunks: {len(documents)}")
    
    # Check if we have any documents to add
    if not documents:
        print("No documents to add to vector store")
        return None
    
    # Create a session-specific collection name
    collection_name = f"session_{session_id}"
    print(f"Collection name: {collection_name}")
    
    # Print first few document contents for debugging
    print(f"First document content preview: {documents[0].page_content[:200]}...")
    print(f"First document metadata: {documents[0].metadata}")
    
    # Create the vector store with explicit persistence
    try:
        print("Initializing Chroma vector store...")
        # Initialize Chroma with the collection name and embedding model
        vector_store = Chroma(
            collection_name=collection_name,
            embedding_function=embedding_model,
            persist_directory="./chroma_db"
        )
        print("Chroma vector store initialized")
        
        # Add documents to the vector store
        print(f"Adding {len(documents)} documents to vector store...")
        print("Starting embedding process...")
        vector_store.add_documents(documents)
        print("Embedding process completed")
        print("Documents added to vector store")
        
        # Verify that documents were added
        print("Checking document count...")
        count = vector_store._collection.count()
        print(f"Vector store created successfully with {count} documents")
        
        # If count is 0, try to force a reload
        if count == 0:
            print("Document count is 0, trying to reload vector store...")
            # Re-initialize the vector store to force a reload
            vector_store = Chroma(
                collection_name=collection_name,
                embedding_function=embedding_model,
                persist_directory="./chroma_db"
            )
            count = vector_store._collection.count()
            print(f"Reloaded vector store with {count} documents")
        
        # List collections in the persist directory for debugging
        try:
            import os
            chroma_db_path = "./chroma_db"
            if os.path.exists(chroma_db_path):
                collections = os.listdir(chroma_db_path)
                print(f"Collections in persist directory: {collections}")
                # Print contents of chroma_db directory
                for collection in collections:
                    collection_path = os.path.join(chroma_db_path, collection)
                    if os.path.isdir(collection_path):
                        collection_files = os.listdir(collection_path)
                        print(f"Files in collection {collection}: {collection_files}")
        except Exception as e:
            print(f"Error listing collections: {str(e)}")
        
        return vector_store
    except Exception as e:
        print(f"Error creating vector store: {str(e)}")
        import traceback
        traceback.print_exc()
        return None

def get_vector_store_by_session(session_id: str) -> Chroma:
    """
    Get a Chroma vector store by session ID.
    
    Args:
        session_id (str): Session ID for collection naming.
        
    Returns:
        Chroma: Vector store with the document chunks.
    """
    # Create a session-specific collection name
    collection_name = f"session_{session_id}"
    
    print(f"Getting vector store for session: {session_id}")
    print(f"Collection name: {collection_name}")
    
    # Check if the collection exists
    try:
        # Try to load the vector store
        vector_store = Chroma(
            collection_name=collection_name,
            embedding_function=embedding_model,
            persist_directory="./chroma_db"  # Explicitly set persist directory
        )
        
        # Verify that the collection exists by trying to get its count
        count = vector_store._collection.count()
        print(f"Vector store loaded successfully with {count} documents")
        
        # Also print some sample document metadata to verify
        if count > 0:
            # Get a sample document to verify the structure
            sample_docs = vector_store._collection.get(limit=1)
            if sample_docs and 'metadatas' in sample_docs and sample_docs['metadatas']:
                print(f"Sample document metadata: {sample_docs['metadatas'][0]}")
        else:
            print("Vector store is empty")
            # List available collections to help with debugging
            try:
                import os
                chroma_db_path = "./chroma_db"
                if os.path.exists(chroma_db_path):
                    collections = os.listdir(chroma_db_path)
                    print(f"Available collections: {collections}")
                    # Print contents of chroma_db directory
                    for collection in collections:
                        collection_path = os.path.join(chroma_db_path, collection)
                        if os.path.isdir(collection_path):
                            collection_files = os.listdir(collection_path)
                            print(f"Files in collection {collection}: {collection_files}")
            except Exception as list_error:
                print(f"Error listing collections: {str(list_error)}")
        
        return vector_store
    except Exception as e:
        print(f"Error loading vector store: {str(e)}")
        # List available collections to help with debugging
        try:
            import os
            chroma_db_path = "./chroma_db"
            if os.path.exists(chroma_db_path):
                collections = os.listdir(chroma_db_path)
                print(f"Available collections: {collections}")
        except Exception as list_error:
            print(f"Error listing collections: {str(list_error)}")
        
        raise Exception(f"Could not load vector store for session {session_id}: {str(e)}")


def cleanup_session_data(session_id: str):
    """
    Clean up Chroma collections for a given session ID.
    
    Args:
        session_id (str): Session ID for which to delete collections.
    """
    import os
    import shutil
    
    # Chroma stores data in a directory named 'chroma_db' by default
    chroma_db_path = "./chroma_db"
    
    if os.path.exists(chroma_db_path):
        # Iterate through subdirectories to find collections for this session
        for item in os.listdir(chroma_db_path):
            item_path = os.path.join(chroma_db_path, item)
            # Check if the item is a directory and matches the session ID pattern
            if os.path.isdir(item_path) and item.startswith(f"session_{session_id}_"):
                try:
                    shutil.rmtree(item_path)
                    print(f"Deleted Chroma collection: {item}")
                except Exception as e:
                    print(f"Error deleting Chroma collection {item}: {str(e)}")
    else:
        print(f"Chroma DB path {chroma_db_path} does not exist")


def retrieve_relevant_chunks(vector_store: Chroma, query: str, k: int = 4, filter: Optional[Dict[str, Any]] = None) -> List[Document]:
    """
    Retrieve relevant chunks from the vector store based on a query.
    Uses a sophisticated retrieval strategy optimized for large context windows.
    
    Args:
        vector_store (Chroma): Vector store to search in.
        query (str): Query to search for.
        k (int): Number of chunks to retrieve (fewer for larger context windows).
        filter (Optional[Dict[str, Any]]): Metadata filter to apply during search.
        
    Returns:
        List[Document]: List of relevant document chunks.
    """
    print(f"Retrieving chunks for query: {query}")
    print(f"Requested number of chunks: {k}")
    if filter:
        print(f"Applying metadata filter: {filter}")
    
    # Add a small delay to ensure embedding process completes
    import time
    time.sleep(0.1)
    
    search_kwargs = {
        "k": k,
        "fetch_k": k * 5,  # Retrieve more candidates for better selection
        "lambda_mult": 0.8  # Balance between relevance and diversity
    }
    if filter:
        search_kwargs["filter"] = filter

    # Use MMR (Maximal Marginal Relevance) for diversity with a higher lambda for relevance
    retriever = vector_store.as_retriever(
        search_type="mmr",
        search_kwargs=search_kwargs
    )
    # Use the new invoke method instead of the deprecated get_relevant_documents
    docs = retriever.invoke(query)
    
    print(f"Retrieved {len(docs)} chunks")
    if docs:
        # Print first chunk content for debugging
        print(f"First chunk content preview: {docs[0].page_content[:200]}...")
        print(f"First chunk metadata: {docs[0].metadata}")
    
    return docs

def process_file(file_path: str, session_id: str, max_chunk_size: int = 10000) -> Dict[str, Any]:
    """
    Process a file based on its extension, creating session-specific vector stores.
    
    Args:
        file_path (str): Path to the file.
        session_id (str): Session ID for collection naming.
        max_chunk_size (int): Maximum size of text to process directly (larger texts will use RAG).
        
    Returns:
        Dict[str, Any]: Dictionary containing file information and extracted content.
    """
    print(f"Processing file: {file_path} for session: {session_id}")
    
    if not os.path.exists(file_path):
        return {"error": "File not found"}
    
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    file_info = {
        "file_path": file_path,
        "file_name": os.path.basename(file_path),
        "file_extension": ext,
        "content": "",
        "vector_store": None,
        "chunk_count": 0,
        "processing_log": [],
        "is_image": False  # Default to False, will be set to True for image files
    }
    
    # Add to processing log
    file_info["processing_log"].append(f"Processing file: {file_path}")
    
    try:
        if ext == ".pdf":
            file_info["processing_log"].append("Extracting text from PDF using hybrid approach...")
            file_info["content"] = extract_text_from_pdf(file_path, use_docling=True)
            file_info["processing_log"].append(f"Extracted {len(file_info['content'])} characters from PDF using hybrid approach")
        elif ext == ".docx":
            file_info["processing_log"].append("Extracting text from DOCX using hybrid approach...")
            file_info["content"] = extract_text_from_docx(file_path, use_docling=True)
            file_info["processing_log"].append(f"Extracted {len(file_info['content'])} characters from DOCX using hybrid approach")
        elif ext in [".txt", ".md", ".html", ".py", ".js"]:
            file_info["processing_log"].append("Extracting text from text file...")
            file_info["content"] = extract_text_from_txt(file_path)
            file_info["processing_log"].append(f"Extracted {len(file_info['content'])} characters from text file")
        elif ext == ".csv":
            file_info["processing_log"].append("Extracting data from CSV...")
            file_info["content"] = extract_text_from_csv(file_path)
            file_info["processing_log"].append(f"Extracted CSV data")
        elif ext == ".xlsx":
            file_info["processing_log"].append("Extracting data from Excel...")
            file_info["content"] = extract_text_from_xlsx(file_path)
            file_info["processing_log"].append(f"Extracted Excel data")
        elif ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"]:
            file_info["processing_log"].append("Processing image file...")
            file_info["content"] = extract_text_from_image(file_path)
            file_info["processing_log"].append(f"Processed image file {file_info['file_name']}")
        else:
            # For unsupported file types, try to read as text
            file_info["processing_log"].append("Extracting text from file (generic)...")
            file_info["content"] = extract_text_from_txt(file_path)
            file_info["processing_log"].append(f"Extracted {len(file_info['content'])} characters from file")
    except Exception as e:
        file_info["error"] = str(e)
        file_info["processing_log"].append(f"Error during extraction: {str(e)}")
        return file_info
    
    print(f"Extracted content length: {len(file_info['content'])}")
    file_info["processing_log"].append(f"Total content length: {len(file_info['content'])}")
    
    # Handle all documents with RAG for consistency, except for image files
    if ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tiff", ".webp"]:
        # For image files, we don't need RAG processing
        file_info["processing_log"].append("Image file processed without RAG")
        file_info["content"] = f"[Image File Processed]\n{file_info['content']}"
        file_info["is_image"] = True  # Flag to indicate this is an image file
        file_info["has_vector_store"] = False  # Explicitly set to False for image files
        file_info["processing_log"].append("Image file processing completed")
    else:
        # Handle all other documents with RAG for consistency
        file_info["processing_log"].append(f"Processing document with RAG ({len(file_info['content'])} chars)")
        file_info["processing_log"].append("Chunking text...")
        chunks = chunk_text(file_info["content"], chunk_size=4000, chunk_overlap=400)
        file_info["chunk_count"] = len(chunks)
        file_info["processing_log"].append(f"Created {len(chunks)} chunks")
        
        file_info["processing_log"].append("Creating vector store with embeddings...")
        vector_store = create_vector_store(chunks, file_info["file_name"], session_id)
        if vector_store:
            file_info["vector_store"] = vector_store
            # For display, we'll show a preview of the content with chunk information
            file_info["content"] = f"[Document with {len(chunks)} chunks]\nPreview:\n{file_info['content'][:1000]}..."
            file_info["processing_log"].append(f"Created vector store with {len(chunks)} chunks")
            file_info["processing_log"].append("Document processing with RAG completed successfully")
        else:
            file_info["processing_log"].append("Failed to create vector store, using direct content")
            file_info["processing_log"].append("Document processing completed with direct content (no RAG)")
    
    return file_info